!/usr/bin/env python3
“””
Academic Translator - Making ALL Knowledge Accessible to ALL Minds

Translates academic papers, research studies, and technical documents into formats
that work for different learning styles, reading levels, and cognitive differences.

Modular architecture allows community contributions for specific accessibility needs.
“””

import requests
from bs4 import BeautifulSoup
import re
import json
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional, Tuple, Union
import time
from pathlib import Path
import importlib
from abc import ABC, abstractmethod
import PyPDF2
import fitz  # PyMuPDF for better PDF extraction
from docx import Document

@dataclass
class AcademicTranslationResult:
original_text: str
plain_english: str
key_findings: List[str]
why_this_matters: List[str]
methodology_simplified: List[str]
questions_to_ask: List[str]
action_items: List[str]
visual_elements: List[str]
confidence_score: float
subject_area: str
reading_level: str
modules_applied: List[str]
source_file: str

class AccessibilityModule(ABC):
“”“Base class for all accessibility modules”””

```
@abstractmethod
def get_name(self) -> str:
    """Return module name"""
    pass

@abstractmethod
def get_description(self) -> str:
    """Return module description"""
    pass

@abstractmethod
def process_text(self, text: str, context: Dict) -> str:
    """Process text according to module's accessibility needs"""
    pass

@abstractmethod
def get_additional_elements(self, text: str, context: Dict) -> Dict[str, List[str]]:
    """Return additional elements this module provides"""
    pass
```

class AcademicTranslator:
def **init**(self):
self.academic_jargon = self.load_academic_jargon()
self.subject_patterns = self.load_subject_patterns()
self.methodology_keywords = self.load_methodology_keywords()
self.loaded_modules = {}
self.available_modules = self.discover_modules()

```
def load_academic_jargon(self) -> Dict[str, Dict[str, str]]:
    """Academic jargon translation dictionary organized by field"""
    return {
        'general_research': {
            'methodology': 'how the study was done',
            'literature review': 'summary of previous research on this topic',
            'hypothesis': 'educated guess about what would happen',
            'null hypothesis': 'assumption that there\'s no real effect',
            'sample size': 'number of people/things studied',
            'control group': 'comparison group that didn\'t get the treatment',
            'experimental group': 'group that got the treatment being tested',
            'placebo': 'fake treatment with no active ingredient',
            'double-blind': 'neither participants nor researchers knew who got real treatment',
            'randomized': 'people were randomly assigned to groups',
            'correlation': 'things that tend to happen together (doesn\'t prove cause)',
            'causation': 'one thing actually causes another',
            'statistical significance': 'result is probably not due to chance',
            'p-value': 'probability the result happened by accident',
            'confidence interval': 'range where the true answer probably lies',
            'peer review': 'other experts checked this research before publication',
            'replication': 'repeating the study to see if results hold up',
            'meta-analysis': 'study that combines results from multiple studies',
        },
        'medical': {
            'clinical trial': 'research study testing treatments on people',
            'randomized controlled trial': 'gold standard study where people are randomly assigned treatments',
            'cohort study': 'following a group of people over time',
            'case-control study': 'comparing people with a condition to those without',
            'systematic review': 'comprehensive summary of all research on a topic',
            'efficacy': 'how well treatment works in ideal conditions',
            'effectiveness': 'how well treatment works in real-world conditions',
            'adverse events': 'bad side effects',
            'contraindication': 'reason not to use this treatment',
            'comorbidity': 'having multiple health conditions at once',
            'prevalence': 'how common a condition is',
            'incidence': 'how many new cases occur in a time period',
            'mortality': 'death rate',
            'morbidity': 'illness rate',
            'biomarker': 'measurable sign of disease or treatment effect',
            'pharmacokinetics': 'how the body processes medication',
            'pharmacodynamics': 'how medication affects the body',
        },
        'psychology': {
            'construct': 'concept being measured (like intelligence or depression)',
            'validity': 'whether a test measures what it claims to measure',
            'reliability': 'whether a test gives consistent results',
            'operational definition': 'exact way researchers define and measure something',
            'confounding variable': 'outside factor that might affect results',
            'cognitive bias': 'systematic error in thinking',
            'effect size': 'how big the difference actually is (practical importance)',
            'standard deviation': 'measure of how spread out the data is',
            'normal distribution': 'bell curve - most people in the middle, few at extremes',
            'outlier': 'unusual result that doesn\'t fit the pattern',
        },
        'education': {
            'pedagogical': 'related to teaching methods',
            'scaffolding': 'providing support that\'s gradually removed as students learn',
            'differentiation': 'adapting teaching for different student needs',
            'formative assessment': 'checking understanding during learning',
            'summative assessment': 'final test of what was learned',
            'metacognition': 'thinking about thinking - awareness of your own learning',
            'zone of proximal development': 'sweet spot between too easy and too hard',
            'intrinsic motivation': 'motivation from internal satisfaction',
            'extrinsic motivation': 'motivation from external rewards',
        },
        'social_science': {
            'qualitative research': 'studying experiences, meanings, and perspectives',
            'quantitative research': 'studying numbers and statistics',
            'ethnography': 'studying culture by observing and participating',
            'phenomenology': 'studying people\'s lived experiences',
            'grounded theory': 'developing theory from data rather than testing existing theory',
            'triangulation': 'using multiple methods to confirm findings',
            'thick description': 'rich, detailed account of what was observed',
            'reflexivity': 'researcher reflecting on how they might bias the study',
        },
        'statistics': {
            'mean': 'average',
            'median': 'middle value when all values are arranged in order',
            'mode': 'most common value',
            'range': 'difference between highest and lowest values',
            'variance': 'measure of how spread out data is',
            'regression': 'method to predict one variable from others',
            'anova': 'test to compare averages between groups',
            't-test': 'test to compare averages between two groups',
            'chi-square': 'test for relationships between categories',
            'power': 'ability of a test to detect a real effect',
            'type i error': 'false positive - finding an effect that isn\'t really there',
            'type ii error': 'false negative - missing a real effect',
        }
    }

def load_subject_patterns(self) -> Dict[str, Dict]:
    """Patterns to identify different academic subjects"""
    return {
        'medical': {
            'keywords': ['patient', 'clinical', 'treatment', 'therapy', 'diagnosis', 'symptoms', 'disease', 'health', 'medical', 'hospital', 'doctor', 'physician', 'nurse', 'medication', 'drug', 'surgery', 'intervention'],
            'journals': ['nejm', 'lancet', 'jama', 'bmj', 'nature medicine', 'cell', 'science translational medicine'],
            'methods': ['randomized controlled trial', 'clinical trial', 'cohort study', 'case-control', 'systematic review', 'meta-analysis']
        },
        'psychology': {
            'keywords': ['behavior', 'cognitive', 'mental', 'psychological', 'emotion', 'personality', 'memory', 'learning', 'perception', 'consciousness', 'therapy', 'depression', 'anxiety'],
            'journals': ['psychological science', 'journal of personality', 'cognitive psychology', 'developmental psychology'],
            'methods': ['survey', 'experiment', 'longitudinal study', 'cross-sectional', 'qualitative interview']
        },
        'education': {
            'keywords': ['student', 'learning', 'teaching', 'classroom', 'curriculum', 'pedagogy', 'instruction', 'academic achievement', 'educational', 'school'],
            'journals': ['journal of educational psychology', 'educational researcher', 'review of educational research'],
            'methods': ['action research', 'case study', 'mixed methods', 'educational intervention']
        },
        'social_science': {
            'keywords': ['social', 'society', 'culture', 'community', 'policy', 'economic', 'political', 'demographic', 'inequality', 'justice'],
            'journals': ['american sociological review', 'social forces', 'sociology of education'],
            'methods': ['ethnography', 'survey research', 'content analysis', 'field study']
        },
        'science': {
            'keywords': ['experiment', 'hypothesis', 'theory', 'model', 'analysis', 'data', 'results', 'conclusion', 'research', 'study'],
            'journals': ['nature', 'science', 'cell', 'pnas', 'journal of biological chemistry'],
            'methods': ['laboratory experiment', 'field study', 'modeling', 'simulation']
        }
    }

def load_methodology_keywords(self) -> Dict[str, str]:
    """Common research methodology terms simplified"""
    return {
        'n =': 'number of people/things studied:',
        'p <': 'probability this happened by chance:',
        'r =': 'strength of relationship:',
        'f =': 'statistical test result:',
        't =': 'statistical test result:',
        'χ² =': 'chi-square test result:',
        'df =': 'degrees of freedom (technical statistical term):',
        'ci =': 'confidence interval (range where true answer likely falls):',
        'm =': 'average:',
        'sd =': 'standard deviation (how spread out the data is):',
    }

def discover_modules(self) -> List[str]:
    """Discover available accessibility modules"""
    modules = []
    modules_dir = Path(__file__).parent / 'modules'
    
    if modules_dir.exists():
        for file in modules_dir.glob('*_module.py'):
            module_name = file.stem
            modules.append(module_name)
    
    return modules

def load_module(self, module_name: str) -> Optional[AccessibilityModule]:
    """Load a specific accessibility module"""
    if module_name in self.loaded_modules:
        return self.loaded_modules[module_name]
    
    try:
        module_path = f"modules.{module_name}"
        module = importlib.import_module(module_path)
        
        # Find the module class (should end with 'Module')
        for attr_name in dir(module):
            attr = getattr(module, attr_name)
            if (isinstance(attr, type) and 
                issubclass(attr, AccessibilityModule) and 
                attr != AccessibilityModule):
                
                instance = attr()
                self.loaded_modules[module_name] = instance
                return instance
    
    except ImportError as e:
        print(f"Could not load module {module_name}: {e}")
    
    return None

def detect_subject_area(self, text: str) -> str:
    """Identify the academic subject area"""
    text_lower = text.lower()
    
    scores = {}
    for subject, patterns in self.subject_patterns.items():
        score = 0
        
        # Check keywords
        for keyword in patterns['keywords']:
            score += text_lower.count(keyword) * 2
        
        # Check journal patterns
        for journal in patterns.get('journals', []):
            if journal.lower() in text_lower:
                score += 5
        
        # Check methodology patterns
        for method in patterns.get('methods', []):
            if method.lower() in text_lower:
                score += 3
        
        scores[subject] = score
    
    return max(scores, key=scores.get) if scores else 'general'

def extract_text_from_file(self, file_path: str) -> str:
    """Extract text from various academic document formats"""
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.pdf':
        return self.extract_text_from_pdf(str(file_path))
    elif file_path.suffix.lower() in ['.docx', '.doc']:
        return self.extract_text_from_docx(str(file_path))
    elif file_path.suffix.lower() in ['.txt', '.text']:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

def extract_text_from_pdf(self, file_path: str) -> str:
    """Extract text from PDF - academic papers often have complex layouts"""
    text = ""
    
    try:
        # Try PyMuPDF first (better for academic papers)
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        
        if len(text.strip()) > 100:
            return text
            
    except Exception as e:
        print(f"PyMuPDF failed: {e}, trying PyPDF2...")
    
    try:
        # Fallback to PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
                
    except Exception as e:
        print(f"PyPDF2 also failed: {e}")
        return ""
    
    return text

def extract_text_from_docx(self, file_path: str) -> str:
    """Extract text from Word documents"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Error reading Word document: {e}")
        return ""

def translate_academic_jargon(self, text: str, subject_area: str) -> str:
    """Replace academic jargon with plain English"""
    translated = text
    
    # Apply general research terms
    for jargon, plain in self.academic_jargon['general_research'].items():
        pattern = r'\b' + re.escape(jargon) + r'\b'
        replacement = f"{plain} ({jargon})"
        translated = re.sub(pattern, replacement, translated, flags=re.IGNORECASE)
    
    # Apply subject-specific terms
    if subject_area in self.academic_jargon:
        for jargon, plain in self.academic_jargon[subject_area].items():
            pattern = r'\b' + re.escape(jargon) + r'\b'
            replacement = f"{plain} ({jargon})"
            translated = re.sub(pattern, replacement, translated, flags=re.IGNORECASE)
    
    # Apply statistical/methodology terms
    for jargon, plain in self.methodology_keywords.items():
        pattern = re.escape(jargon)
        translated = re.sub(pattern, plain, translated, flags=re.IGNORECASE)
    
    return translated

def extract_key_findings(self, text: str, subject_area: str) -> List[str]:
    """Extract the main research findings"""
    findings = []
    
    # Look for results/findings sections
    results_patterns = [
        r'(?:results?|findings?|outcomes?)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:we found|our findings|the study found|results showed)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:conclusion|conclusions)[:\s]([^.]*(?:\.[^.]*){0,2})'
    ]
    
    for pattern in results_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            finding = match.group(1).strip()
            if len(finding) > 20:  # Filter out very short matches
                findings.append(f"🔍 {finding}")
    
    # Look for statistically significant results
    sig_patterns = [
        r'(p\s*[<>=]\s*0\.0[0-5][^.]*)',
        r'(significant[ly]?\s+[^.]*)',
        r'([^.]*significant\s+(?:difference|effect|relationship)[^.]*)'
    ]
    
    for pattern in sig_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            result = match.group(1).strip()
            findings.append(f"📊 {result}")
    
    return findings[:8]  # Limit to most important

def extract_methodology_simplified(self, text: str, subject_area: str) -> List[str]:
    """Simplify the research methods section"""
    methods = []
    
    # Look for methods section
    methods_patterns = [
        r'(?:methods?|methodology|procedure|design)[:\s]([^.]*(?:\.[^.]*){0,3})',
        r'(?:participants?|subjects?)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:data collection|analysis)[:\s]([^.]*(?:\.[^.]*){0,2})'
    ]
    
    for pattern in methods_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            method = match.group(1).strip()
            if len(method) > 15:
                methods.append(f"⚙️ {method}")
    
    # Look for sample size information
    sample_patterns = [
        r'(n\s*=\s*\d+[^.]*)',
        r'(\d+\s+(?:participants?|subjects?|patients?)[^.]*)',
        r'(sample\s+(?:of|size)[^.]*\d+[^.]*)'
    ]
    
    for pattern in sample_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            sample_info = match.group(1).strip()
            methods.append(f"👥 {sample_info}")
    
    return methods[:6]

def generate_why_this_matters(self, text: str, subject_area: str) -> List[str]:
    """Generate "why this matters" explanations"""
    matters = []
    
    # Subject-specific implications
    if subject_area == 'medical':
        matters.extend([
            "💊 Could lead to better treatments for patients",
            "🏥 May change how doctors diagnose or treat conditions",
            "⚠️ Might reveal new risks or benefits of treatments",
            "🔬 Advances our understanding of how the body works"
        ])
    elif subject_area == 'psychology':
        matters.extend([
            "🧠 Helps us understand how the mind works",
            "🤝 Could improve relationships and social interactions", 
            "📚 May change how we approach learning and education",
            "💭 Provides insights into human behavior and decision-making"
        ])
    elif subject_area == 'education':
        matters.extend([
            "🎓 Could improve how students learn",
            "👩‍🏫 May help teachers be more effective",
            "📈 Might boost academic achievement",
            "🌍 Could reduce educational inequalities"
        ])
    
    # Look for implications/discussion sections
    impl_patterns = [
        r'(?:implications?|significance|impact)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:these findings|our results|this study)\s+(?:suggest|indicate|show)[s]?\s+([^.]*)',
        r'(?:clinical|practical|policy)\s+implications[:\s]([^.]*)'
    ]
    
    for pattern in impl_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            implication = match.group(1).strip()
            if len(implication) > 20:
                matters.append(f"🎯 {implication}")
    
    return matters[:6]

def generate_questions_to_ask(self, text: str, subject_area: str) -> List[str]:
    """Generate questions people should ask professionals"""
    questions = []
    
    if subject_area == 'medical':
        questions.extend([
            "🩺 Ask your doctor: 'Does this research apply to my specific condition?'",
            "💊 'Should I change my current treatment based on this?'",
            "⚠️ 'What are the risks and benefits for someone like me?'",
            "🔍 'Are there clinical trials I could participate in?'"
        ])
    elif subject_area == 'psychology':
        questions.extend([
            "🧠 Ask a therapist: 'How might this apply to my situation?'",
            "📚 'Could this research help me understand my behavior better?'",
            "🤝 'How can I use this information in my relationships?'"
        ])
    elif subject_area == 'education':
        questions.extend([
            "👩‍🏫 Ask teachers: 'How could this improve my child's learning?'",
            "📖 'Should we change our study methods based on this?'",
            "🎓 'What does this mean for educational policy?'"
        ])
    
    # Look for limitations mentioned in the study
    limitations = re.findall(r'limitation[s]?\s+(?:of this study\s+)?(?:include|are)[:\s]([^.]*)', text, re.IGNORECASE)
    for limitation in limitations:
        questions.append(f"❓ Ask experts: 'How do the study limitations affect the conclusions?'")
    
    return questions[:6]

def calculate_reading_level(self, text: str) -> str:
    """Estimate reading level of the text"""
    # Simple reading level estimation based on sentence length and word complexity
    sentences = re.split(r'[.!?]+', text)
    words = text.split()
    
    if not sentences or not words:
        return "Unknown"
    
    avg_sentence_length = len(words) / len(sentences)
    
    # Count complex words (3+ syllables, rough estimate)
    complex_words = 0
    for word in words[:1000]:  # Sample first 1000 words
        if len(word) > 6:  # Rough proxy for syllable count
            complex_words += 1
    
    complex_ratio = complex_words / min(len(words), 1000)
    
    # Very rough reading level estimation
    if avg_sentence_length < 15 and complex_ratio < 0.1:
        return "Middle School"
    elif avg_sentence_length < 20 and complex_ratio < 0.15:
        return "High School"
    elif avg_sentence_length < 25 and complex_ratio < 0.2:
        return "College"
    else:
        return "Graduate/Professional"

def translate_academic_document(self, text: str, modules: List[str] = None, 
                              source_file: str = None) -> AcademicTranslationResult:
    """Main translation function with modular accessibility support"""
    
    # Detect subject area
    subject_area = self.detect_subject_area(text)
    reading_level = self.calculate_reading_level(text)
    
    # Apply core translation
    plain_text = self.translate_academic_jargon(text, subject_area)
    
    # Extract core information
    key_findings = self.extract_key_findings(text, subject_area)
    methodology = self.extract_methodology_simplified(text, subject_area)
    why_matters = self.generate_why_this_matters(text, subject_area)
    questions = self.generate_questions_to_ask(text, subject_area)
    
    # Apply accessibility modules
    applied_modules = []
    additional_elements = {'visual_elements': [], 'action_items': []}
    
    if modules:
        for module_name in modules:
            module = self.load_module(module_name)
            if module:
                context = {
                    'subject_area': subject_area,
                    'reading_level': reading_level,
                    'key_findings': key_findings
                }
                
                # Apply module transformations
                plain_text = module.process_text(plain_text, context)
                
                # Get additional elements
                module_elements = module.get_additional_elements(text, context)
                for key, values in module_elements.items():
                    if key in additional_elements:
                        additional_elements[key].extend(values)
                
                applied_modules.append(module.get_name())
    
    confidence = self.calculate_confidence(text, subject_area, applied_modules)
    
    return AcademicTranslationResult(
        original_text=text,
        plain_english=plain_text,
        key_findings=key_findings,
        why_this_matters=why_matters,
        methodology_simplified=methodology,
        questions_to_ask=questions,
        action_items=additional_elements.get('action_items', []),
        visual_elements=additional_elements.get('visual_elements', []),
        confidence_score=confidence,
        subject_area=subject_area,
        reading_level=reading_level,
        modules_applied=applied_modules,
        source_file=source_file or "Direct input"
    )

def calculate_confidence(self, text: str, subject_area: str, modules: List[str]) -> float:
    """Calculate confidence in translation quality"""
    base_score = 0.7
    
    # Boost for recognized subject area
    if subject_area != 'general':
        base_score += 0.1
    
    # Boost for finding key research elements
    if re.search(r'(?:results?|findings?|conclusion)', text, re.IGNORECASE):
        base_score += 0.1
    if re.search(r'(?:methods?|methodology)', text, re.IGNORECASE):
        base_score += 0.1
    
    # Boost for applied modules
    base_score += len(modules) * 0.02
    
    # Penalize very short or very long texts
    if len(text) < 1000:
        base_score -= 0.1
    elif len(text) > 50000:
        base_score -= 0.1
    
    return max(0.3, min(0.95, base_score))

def save_translation(self, result: AcademicTranslationResult, filename: str):
    """Save translation results"""
    Path("academic_translations").mkdir(exist_ok=True)
    
    # Save JSON
    with open(f"academic_translations/{filename}.json", 'w') as f:
        json.dump(asdict(result), f, indent=2)
    
    # Save HTML report
    html_report = self.generate_html_report(result)
    with open(f"academic_translations/{filename}.html", 'w') as f:
        f.write(html_report)

def generate_html_report(self, result: AcademicTranslationResult) -> str:
    """Generate comprehensive HTML report"""
    html = f"""
```

<!DOCTYPE html>

<html>
<head>
    <title>Academic Translation: {result.subject_area.title()}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 1000px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 10px; margin-bottom: 30px; }}
        .subject {{ font-size: 28px; font-weight: bold; }}
        .reading-level {{ font-size: 16px; opacity: 0.9; margin-top: 10px; }}
        .modules {{ font-size: 14px; margin-top: 15px; }}
        .section {{ margin: 25px 0; padding: 20px; border-left: 5px solid #667eea; background: #f8faff; border-radius: 5px; }}
        .findings {{ border-left-color: #4CAF50; background: #f8fff8; }}
        .methodology {{ border-left-color: #FF9800; background: #fff8f0; }}
        .matters {{ border-left-color: #E91E63; background: #fdf8fb; }}
        .questions {{ border-left-color: #9C27B0; background: #faf8ff; }}
        .confidence {{ text-align: center; font-size: 20px; margin: 30px 0; padding: 15px; background: #e3f2fd; border-radius: 8px; }}
        ul {{ padding-left: 20px; }}
        li {{ margin: 10px 0; }}
        .key-point {{ font-weight: bold; color: #1976D2; }}
        h2 {{ color: #333; border-bottom: 2px solid #eee; padding-bottom: 10px; }}
        .visual-note {{ background: #fff3e0; padding: 15px; border-radius: 5px; margin: 10px 0; border-left: 4px solid #ff9800; }}
    </style>
</head>
<body>
    <div class="header">
        <div class="subject">📚 {result.subject_area.title()} Research Translation</div>
        <div class="reading-level">📖 Original Reading Level: {result.reading_level}</div>
        <div class="modules">🔧 Accessibility Modules Applied: {', '.join(result.modules_applied) if result.modules_applied else 'None'}</div>
    </div>

```
<div class="confidence">
    <strong>🎯 Translation Confidence: {result.confidence_score:.0%}</strong>
</div>

<div class="section findings">
```

#!/usr/bin/env python3
“””
Academic Translator - Making ALL Knowledge Accessible to ALL Minds

Translates academic papers, research studies, and technical documents into formats
that work for different learning styles, reading levels, and cognitive differences.

Modular architecture allows community contributions for specific accessibility needs.
“””

import requests
from bs4 import BeautifulSoup
import re
import json
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional, Tuple, Union
import time
from pathlib import Path
import importlib
from abc import ABC, abstractmethod
import PyPDF2
import fitz  # PyMuPDF for better PDF extraction
from docx import Document

@dataclass
class AcademicTranslationResult:
original_text: str
plain_english: str
key_findings: List[str]
why_this_matters: List[str]
methodology_simplified: List[str]
questions_to_ask: List[str]
action_items: List[str]
visual_elements: List[str]
confidence_score: float
subject_area: str
reading_level: str
modules_applied: List[str]
source_file: str

class AccessibilityModule(ABC):
“”“Base class for all accessibility modules”””

```
@abstractmethod
def get_name(self) -> str:
    """Return module name"""
    pass

@abstractmethod
def get_description(self) -> str:
    """Return module description"""
    pass

@abstractmethod
def process_text(self, text: str, context: Dict) -> str:
    """Process text according to module's accessibility needs"""
    pass

@abstractmethod
def get_additional_elements(self, text: str, context: Dict) -> Dict[str, List[str]]:
    """Return additional elements this module provides"""
    pass
```

class AcademicTranslator:
def **init**(self):
self.academic_jargon = self.load_academic_jargon()
self.subject_patterns = self.load_subject_patterns()
self.methodology_keywords = self.load_methodology_keywords()
self.loaded_modules = {}
self.available_modules = self.discover_modules()

```
def load_academic_jargon(self) -> Dict[str, Dict[str, str]]:
    """Academic jargon translation dictionary organized by field"""
    return {
        'general_research': {
            'methodology': 'how the study was done',
            'literature review': 'summary of previous research on this topic',
            'hypothesis': 'educated guess about what would happen',
            'null hypothesis': 'assumption that there\'s no real effect',
            'sample size': 'number of people/things studied',
            'control group': 'comparison group that didn\'t get the treatment',
            'experimental group': 'group that got the treatment being tested',
            'placebo': 'fake treatment with no active ingredient',
            'double-blind': 'neither participants nor researchers knew who got real treatment',
            'randomized': 'people were randomly assigned to groups',
            'correlation': 'things that tend to happen together (doesn\'t prove cause)',
            'causation': 'one thing actually causes another',
            'statistical significance': 'result is probably not due to chance',
            'p-value': 'probability the result happened by accident',
            'confidence interval': 'range where the true answer probably lies',
            'peer review': 'other experts checked this research before publication',
            'replication': 'repeating the study to see if results hold up',
            'meta-analysis': 'study that combines results from multiple studies',
        },
        'medical': {
            'clinical trial': 'research study testing treatments on people',
            'randomized controlled trial': 'gold standard study where people are randomly assigned treatments',
            'cohort study': 'following a group of people over time',
            'case-control study': 'comparing people with a condition to those without',
            'systematic review': 'comprehensive summary of all research on a topic',
            'efficacy': 'how well treatment works in ideal conditions',
            'effectiveness': 'how well treatment works in real-world conditions',
            'adverse events': 'bad side effects',
            'contraindication': 'reason not to use this treatment',
            'comorbidity': 'having multiple health conditions at once',
            'prevalence': 'how common a condition is',
            'incidence': 'how many new cases occur in a time period',
            'mortality': 'death rate',
            'morbidity': 'illness rate',
            'biomarker': 'measurable sign of disease or treatment effect',
            'pharmacokinetics': 'how the body processes medication',
            'pharmacodynamics': 'how medication affects the body',
        },
        'psychology': {
            'construct': 'concept being measured (like intelligence or depression)',
            'validity': 'whether a test measures what it claims to measure',
            'reliability': 'whether a test gives consistent results',
            'operational definition': 'exact way researchers define and measure something',
            'confounding variable': 'outside factor that might affect results',
            'cognitive bias': 'systematic error in thinking',
            'effect size': 'how big the difference actually is (practical importance)',
            'standard deviation': 'measure of how spread out the data is',
            'normal distribution': 'bell curve - most people in the middle, few at extremes',
            'outlier': 'unusual result that doesn\'t fit the pattern',
        },
        'education': {
            'pedagogical': 'related to teaching methods',
            'scaffolding': 'providing support that\'s gradually removed as students learn',
            'differentiation': 'adapting teaching for different student needs',
            'formative assessment': 'checking understanding during learning',
            'summative assessment': 'final test of what was learned',
            'metacognition': 'thinking about thinking - awareness of your own learning',
            'zone of proximal development': 'sweet spot between too easy and too hard',
            'intrinsic motivation': 'motivation from internal satisfaction',
            'extrinsic motivation': 'motivation from external rewards',
        },
        'social_science': {
            'qualitative research': 'studying experiences, meanings, and perspectives',
            'quantitative research': 'studying numbers and statistics',
            'ethnography': 'studying culture by observing and participating',
            'phenomenology': 'studying people\'s lived experiences',
            'grounded theory': 'developing theory from data rather than testing existing theory',
            'triangulation': 'using multiple methods to confirm findings',
            'thick description': 'rich, detailed account of what was observed',
            'reflexivity': 'researcher reflecting on how they might bias the study',
        },
        'statistics': {
            'mean': 'average',
            'median': 'middle value when all values are arranged in order',
            'mode': 'most common value',
            'range': 'difference between highest and lowest values',
            'variance': 'measure of how spread out data is',
            'regression': 'method to predict one variable from others',
            'anova': 'test to compare averages between groups',
            't-test': 'test to compare averages between two groups',
            'chi-square': 'test for relationships between categories',
            'power': 'ability of a test to detect a real effect',
            'type i error': 'false positive - finding an effect that isn\'t really there',
            'type ii error': 'false negative - missing a real effect',
        }
    }

def load_subject_patterns(self) -> Dict[str, Dict]:
    """Patterns to identify different academic subjects"""
    return {
        'medical': {
            'keywords': ['patient', 'clinical', 'treatment', 'therapy', 'diagnosis', 'symptoms', 'disease', 'health', 'medical', 'hospital', 'doctor', 'physician', 'nurse', 'medication', 'drug', 'surgery', 'intervention'],
            'journals': ['nejm', 'lancet', 'jama', 'bmj', 'nature medicine', 'cell', 'science translational medicine'],
            'methods': ['randomized controlled trial', 'clinical trial', 'cohort study', 'case-control', 'systematic review', 'meta-analysis']
        },
        'psychology': {
            'keywords': ['behavior', 'cognitive', 'mental', 'psychological', 'emotion', 'personality', 'memory', 'learning', 'perception', 'consciousness', 'therapy', 'depression', 'anxiety'],
            'journals': ['psychological science', 'journal of personality', 'cognitive psychology', 'developmental psychology'],
            'methods': ['survey', 'experiment', 'longitudinal study', 'cross-sectional', 'qualitative interview']
        },
        'education': {
            'keywords': ['student', 'learning', 'teaching', 'classroom', 'curriculum', 'pedagogy', 'instruction', 'academic achievement', 'educational', 'school'],
            'journals': ['journal of educational psychology', 'educational researcher', 'review of educational research'],
            'methods': ['action research', 'case study', 'mixed methods', 'educational intervention']
        },
        'social_science': {
            'keywords': ['social', 'society', 'culture', 'community', 'policy', 'economic', 'political', 'demographic', 'inequality', 'justice'],
            'journals': ['american sociological review', 'social forces', 'sociology of education'],
            'methods': ['ethnography', 'survey research', 'content analysis', 'field study']
        },
        'science': {
            'keywords': ['experiment', 'hypothesis', 'theory', 'model', 'analysis', 'data', 'results', 'conclusion', 'research', 'study'],
            'journals': ['nature', 'science', 'cell', 'pnas', 'journal of biological chemistry'],
            'methods': ['laboratory experiment', 'field study', 'modeling', 'simulation']
        }
    }

def load_methodology_keywords(self) -> Dict[str, str]:
    """Common research methodology terms simplified"""
    return {
        'n =': 'number of people/things studied:',
        'p <': 'probability this happened by chance:',
        'r =': 'strength of relationship:',
        'f =': 'statistical test result:',
        't =': 'statistical test result:',
        'χ² =': 'chi-square test result:',
        'df =': 'degrees of freedom (technical statistical term):',
        'ci =': 'confidence interval (range where true answer likely falls):',
        'm =': 'average:',
        'sd =': 'standard deviation (how spread out the data is):',
    }

def discover_modules(self) -> List[str]:
    """Discover available accessibility modules"""
    modules = []
    modules_dir = Path(__file__).parent / 'modules'
    
    if modules_dir.exists():
        for file in modules_dir.glob('*_module.py'):
            module_name = file.stem
            modules.append(module_name)
    
    return modules

def load_module(self, module_name: str) -> Optional[AccessibilityModule]:
    """Load a specific accessibility module"""
    if module_name in self.loaded_modules:
        return self.loaded_modules[module_name]
    
    try:
        module_path = f"modules.{module_name}"
        module = importlib.import_module(module_path)
        
        # Find the module class (should end with 'Module')
        for attr_name in dir(module):
            attr = getattr(module, attr_name)
            if (isinstance(attr, type) and 
                issubclass(attr, AccessibilityModule) and 
                attr != AccessibilityModule):
                
                instance = attr()
                self.loaded_modules[module_name] = instance
                return instance
    
    except ImportError as e:
        print(f"Could not load module {module_name}: {e}")
    
    return None

def detect_subject_area(self, text: str) -> str:
    """Identify the academic subject area"""
    text_lower = text.lower()
    
    scores = {}
    for subject, patterns in self.subject_patterns.items():
        score = 0
        
        # Check keywords
        for keyword in patterns['keywords']:
            score += text_lower.count(keyword) * 2
        
        # Check journal patterns
        for journal in patterns.get('journals', []):
            if journal.lower() in text_lower:
                score += 5
        
        # Check methodology patterns
        for method in patterns.get('methods', []):
            if method.lower() in text_lower:
                score += 3
        
        scores[subject] = score
    
    return max(scores, key=scores.get) if scores else 'general'

def extract_text_from_file(self, file_path: str) -> str:
    """Extract text from various academic document formats"""
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.pdf':
        return self.extract_text_from_pdf(str(file_path))
    elif file_path.suffix.lower() in ['.docx', '.doc']:
        return self.extract_text_from_docx(str(file_path))
    elif file_path.suffix.lower() in ['.txt', '.text']:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

def extract_text_from_pdf(self, file_path: str) -> str:
    """Extract text from PDF - academic papers often have complex layouts"""
    text = ""
    
    try:
        # Try PyMuPDF first (better for academic papers)
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        
        if len(text.strip()) > 100:
            return text
            
    except Exception as e:
        print(f"PyMuPDF failed: {e}, trying PyPDF2...")
    
    try:
        # Fallback to PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
                
    except Exception as e:
        print(f"PyPDF2 also failed: {e}")
        return ""
    
    return text

def extract_text_from_docx(self, file_path: str) -> str:
    """Extract text from Word documents"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Error reading Word document: {e}")
        return ""

def translate_academic_jargon(self, text: str, subject_area: str) -> str:
    """Replace academic jargon with plain English"""
    translated = text
    
    # Apply general research terms
    for jargon, plain in self.academic_jargon['general_research'].items():
        pattern = r'\b' + re.escape(jargon) + r'\b'
        replacement = f"{plain} ({jargon})"
        translated = re.sub(pattern, replacement, translated, flags=re.IGNORECASE)
    
    # Apply subject-specific terms
    if subject_area in self.academic_jargon:
        for jargon, plain in self.academic_jargon[subject_area].items():
            pattern = r'\b' + re.escape(jargon) + r'\b'
            replacement = f"{plain} ({jargon})"
            translated = re.sub(pattern, replacement, translated, flags=re.IGNORECASE)
    
    # Apply statistical/methodology terms
    for jargon, plain in self.methodology_keywords.items():
        pattern = re.escape(jargon)
        translated = re.sub(pattern, plain, translated, flags=re.IGNORECASE)
    
    return translated

def extract_key_findings(self, text: str, subject_area: str) -> List[str]:
    """Extract the main research findings"""
    findings = []
    
    # Look for results/findings sections
    results_patterns = [
        r'(?:results?|findings?|outcomes?)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:we found|our findings|the study found|results showed)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:conclusion|conclusions)[:\s]([^.]*(?:\.[^.]*){0,2})'
    ]
    
    for pattern in results_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            finding = match.group(1).strip()
            if len(finding) > 20:  # Filter out very short matches
                findings.append(f"🔍 {finding}")
    
    # Look for statistically significant results
    sig_patterns = [
        r'(p\s*[<>=]\s*0\.0[0-5][^.]*)',
        r'(significant[ly]?\s+[^.]*)',
        r'([^.]*significant\s+(?:difference|effect|relationship)[^.]*)'
    ]
    
    for pattern in sig_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            result = match.group(1).strip()
            findings.append(f"📊 {result}")
    
    return findings[:8]  # Limit to most important

def extract_methodology_simplified(self, text: str, subject_area: str) -> List[str]:
    """Simplify the research methods section"""
    methods = []
    
    # Look for methods section
    methods_patterns = [
        r'(?:methods?|methodology|procedure|design)[:\s]([^.]*(?:\.[^.]*){0,3})',
        r'(?:participants?|subjects?)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:data collection|analysis)[:\s]([^.]*(?:\.[^.]*){0,2})'
    ]
    
    for pattern in methods_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            method = match.group(1).strip()
            if len(method) > 15:
                methods.append(f"⚙️ {method}")
    
    # Look for sample size information
    sample_patterns = [
        r'(n\s*=\s*\d+[^.]*)',
        r'(\d+\s+(?:participants?|subjects?|patients?)[^.]*)',
        r'(sample\s+(?:of|size)[^.]*\d+[^.]*)'
    ]
    
    for pattern in sample_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            sample_info = match.group(1).strip()
            methods.append(f"👥 {sample_info}")
    
    return methods[:6]

def generate_why_this_matters(self, text: str, subject_area: str) -> List[str]:
    """Generate "why this matters" explanations"""
    matters = []
    
    # Subject-specific implications
    if subject_area == 'medical':
        matters.extend([
            "💊 Could lead to better treatments for patients",
            "🏥 May change how doctors diagnose or treat conditions",
            "⚠️ Might reveal new risks or benefits of treatments",
            "🔬 Advances our understanding of how the body works"
        ])
    elif subject_area == 'psychology':
        matters.extend([
            "🧠 Helps us understand how the mind works",
            "🤝 Could improve relationships and social interactions", 
            "📚 May change how we approach learning and education",
            "💭 Provides insights into human behavior and decision-making"
        ])
    elif subject_area == 'education':
        matters.extend([
            "🎓 Could improve how students learn",
            "👩‍🏫 May help teachers be more effective",
            "📈 Might boost academic achievement",
            "🌍 Could reduce educational inequalities"
        ])
    
    # Look for implications/discussion sections
    impl_patterns = [
        r'(?:implications?|significance|impact)[:\s]([^.]*(?:\.[^.]*){0,2})',
        r'(?:these findings|our results|this study)\s+(?:suggest|indicate|show)[s]?\s+([^.]*)',
        r'(?:clinical|practical|policy)\s+implications[:\s]([^.]*)'
    ]
    
    for pattern in impl_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            implication = match.group(1).strip()
            if len(implication) > 20:
                matters.append(f"🎯 {implication}")
    
    return matters[:6]

def generate_questions_to_ask(self, text: str, subject_area: str) -> List[str]:
    """Generate questions people should ask professionals"""
    questions = []
    
    if subject_area == 'medical':
        questions.extend([
            "🩺 Ask your doctor: 'Does this research apply to my specific condition?'",
            "💊 'Should I change my current treatment based on this?'",
            "⚠️ 'What are the risks and benefits for someone like me?'",
            "🔍 'Are there clinical trials I could participate in?'"
        ])
    elif subject_area == 'psychology':
        questions.extend([
            "🧠 Ask a therapist: 'How might this apply to my situation?'",
            "📚 'Could this research help me understand my behavior better?'",
            "🤝 'How can I use this information in my relationships?'"
        ])
    elif subject_area == 'education':
        questions.extend([
            "👩‍🏫 Ask teachers: 'How could this improve my child's learning?'",
            "📖 'Should we change our study methods based on this?'",
            "🎓 'What does this mean for educational policy?'"
        ])
    
    # Look for limitations mentioned in the study
    limitations = re.findall(r'limitation[s]?\s+(?:of this study\s+)?(?:include|are)[:\s]([^.]*)', text, re.IGNORECASE)
    for limitation in limitations:
        questions.append(f"❓ Ask experts: 'How do the study limitations affect the conclusions?'")
    
    return questions[:6]

def calculate_reading_level(self, text: str) -> str:
    """Estimate reading level of the text"""
    # Simple reading level estimation based on sentence length and word complexity
    sentences = re.split(r'[.!?]+', text)
    words = text.split()
    
    if not sentences or not words:
        return "Unknown"
    
    avg_sentence_length = len(words) / len(sentences)
    
    # Count complex words (3+ syllables, rough estimate)
    complex_words = 0
    for word in words[:1000]:  # Sample first 1000 words
        if len(word) > 6:  # Rough proxy for syllable count
            complex_words += 1
    
    complex_ratio = complex_words / min(len(words), 1000)
    
    # Very rough reading level estimation
    if avg_sentence_length < 15 and complex_ratio < 0.1:
        return "Middle School"
    elif avg_sentence_length < 20 and complex_ratio < 0.15:
        return "High School"
    elif avg_sentence_length < 25 and complex_ratio < 0.2:
        return "College"
    else:
        return "Graduate/Professional"

def translate_academic_document(self, text: str, modules: List[str] = None, 
                              source_file: str = None) -> AcademicTranslationResult:
    """Main translation function with modular accessibility support"""
    
    # Detect subject area
    subject_area = self.detect_subject_area(text)
    reading_level = self.calculate_reading_level(text)
    
    # Apply core translation
    plain_text = self.translate_academic_jargon(text, subject_area)
    
    # Extract core information
    key_findings = self.extract_key_findings(text, subject_area)
    methodology = self.extract_methodology_simplified(text, subject_area)
    why_matters = self.generate_why_this_matters(text, subject_area)
    questions = self.generate_questions_to_ask(text, subject_area)
    
    # Apply accessibility modules
    applied_modules = []
    additional_elements = {'visual_elements': [], 'action_items': []}
    
    if modules:
        for module_name in modules:
            module = self.load_module(module_name)
            if module:
                context = {
                    'subject_area': subject_area,
                    'reading_level': reading_level,
                    'key_findings': key_findings
                }
                
                # Apply module transformations
                plain_text = module.process_text(plain_text, context)
                
                # Get additional elements
                module_elements = module.get_additional_elements(text, context)
                for key, values in module_elements.items():
                    if key in additional_elements:
                        additional_elements[key].extend(values)
                
                applied_modules.append(module.get_name())
    
    confidence = self.calculate_confidence(text, subject_area, applied_modules)
    
    return AcademicTranslationResult(
        original_text=text,
        plain_english=plain_text,
        key_findings=key_findings,
        why_this_matters=why_matters,
        methodology_simplified=methodology,
        questions_to_ask=questions,
        action_items=additional_elements.get('action_items', []),
        visual_elements=additional_elements.get('visual_elements', []),
        confidence_score=confidence,
        subject_area=subject_area,
        reading_level=reading_level,
        modules_applied=applied_modules,
        source_file=source_file or "Direct input"
    )

def calculate_confidence(self, text: str, subject_area: str, modules: List[str]) -> float:
    """Calculate confidence in translation quality"""
    base_score = 0.7
    
    # Boost for recognized subject area
    if subject_area != 'general':
        base_score += 0.1
    
    # Boost for finding key research elements
    if re.search(r'(?:results?|findings?|conclusion)', text, re.IGNORECASE):
        base_score += 0.1
    if re.search(r'(?:methods?|methodology)', text, re.IGNORECASE):
        base_score += 0.1
    
    # Boost for applied modules
    base_score += len(modules) * 0.02
    
    # Penalize very short or very long texts
    if len(text) < 1000:
        base_score -= 0.1
    elif len(text) > 50000:
        base_score -= 0.1
    
    return max(0.3, min(0.95, base_score))

def save_translation(self, result: AcademicTranslationResult, filename: str):
    """Save translation results"""
    Path("academic_translations").mkdir(exist_ok=True)
    
    # Save JSON
    with open(f"academic_translations/{filename}.json", 'w') as f:
        json.dump(asdict(result), f, indent=2)
    
    # Save HTML report
    html_report = self.generate_html_report(result)
    with open(f"academic_translations/{filename}.html", 'w') as f:
        f.write(html_report)

def generate_html_report(self, result: AcademicTranslationResult) -> str:
    """Generate comprehensive HTML report"""
    html = f"""
```

<!DOCTYPE html>

<html>
<head>
    <title>Academic Translation: {result.subject_area.title()}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 1000px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 10px; margin-bottom: 30px; }}
        .subject {{ font-size: 28px; font-weight: bold; }}
        .reading-level {{ font-size: 16px; opacity: 0.9; margin-top: 10px; }}
        .modules {{ font-size: 14px; margin-top: 15px; }}
        .section {{ margin: 25px 0; padding: 20px; border-left: 5px solid #667eea; background: #f8faff; border-radius: 5px; }}
        .findings {{ border-left-color: #4CAF50; background: #f8fff8; }}
        .methodology {{ border-left-color: #FF9800; background: #fff8f0; }}
        .matters {{ border-left-color: #E91E63; background: #fdf8fb; }}
        .questions {{ border-left-color: #9C27B0; background: #faf8ff; }}
        .confidence {{ text-align: center; font-size: 20px; margin: 30px 0; padding: 15px; background: #e3f2fd; border-radius: 8px; }}
        ul {{ padding-left: 20px; }}
        li {{ margin: 10px 0; }}
        .key-point {{ font-weight: bold; color: #1976D2; }}
        h2 {{ color: #333; border-bottom: 2px solid #eee; padding-bottom: 10px; }}
        .visual-note {{ background: #fff3e0; padding: 15px; border-radius: 5px; margin: 10px 0; border-left: 4px solid #ff9800; }}
    </style>
</head>
<body>
    <div class="header">
        <div class="subject">📚 {result.subject_area.title()} Research Translation</div>
        <div class="reading-level">📖 Original Reading Level: {result.reading_level}</div>
        <div class="modules">🔧 Accessibility Modules Applied: {', '.join(result.modules_applied) if result.modules_applied else 'None'}</div>
    </div>

```
<div class="confidence">
    <strong>🎯 Translation Confidence: {result.confidence_score:.0%}</strong>
</div>

<div class="section findings">
    <h2>🔍 Key Findings (What They Discovered)</h2>
    <ul>
    {"".join(f"<li>{finding}</li>" for finding in result.key_findings)}
    </ul>
</div>

<div class="section matters">
    <h2>🎯 Why This Matters (Real-World Impact)</h2>
    <ul>
    {"".join(f"<li>{matter}</li>" for matter in result.why_this_matters)}
    </ul>
</div>

<div class="section methodology">
    <h2>⚙️ How They Did It (Methods Simplified)</h2>
    <ul>
    {"".join(f"<li>{method}</li>" for method in result.methodology_simplified)}
    </ul>
</div>

<div class="section questions">
    <h2>❓ Questions to Ask Professionals</h2>
    <ul>
    {"".join(f"<li>{question}</li>" for question in result.questions_to_ask)}
    </ul>
</div>

{f'''
<div class="section">
    <h2>🎬 Visual Elements</h2>
    <ul>
    {"".join(f"<li>{element}</li>" for element in result.visual_elements)}
    </ul>
</div>
''' if result.visual_elements else ''}

{f'''
<div class="section">
    <h2>📋 Action Items</h2>
    <ul>
    {"".join(f"<li>{action}</li>" for action in result.action_items)}
    </ul>
</div>
''' if result.action_items else ''}

<div class="section">
    <h2>📖 Full Translation</h2>
    <div class="visual-note">
        <strong>💡 Reading Tip:</strong> Technical terms are explained in parentheses. Look for colored highlights if you're using accessibility modules.
    </div>
    <p style="text-align: justify; line-height: 1.8;">
    {result.plain_english.replace(chr(10), '</p><p style="text-align: justify; line-height: 1.8;">')}
    </p>
</div>

<hr style="margin: 40px 0;">
<div style="text-align: center; color: #666; font-size: 14px;">
    <p><strong>Academic Translator</strong> - Making knowledge accessible to all minds</p>
    <p>Source: {result.source_file}</p>
    <p><em>This translation provides information, not professional advice. Always consult experts for important decisions.</em></p>
</div>
```

</body>
</html>
        """
        return html

def main():
“”“Command line interface for academic translation”””
import argparse

```
parser = argparse.ArgumentParser(description='Translate academic papers into accessible formats')
parser.add_argument('--file', '-f', help='Academic paper file (PDF, DOCX, TXT)')
parser.add_argument('--text', '-t', help='Direct text input')
parser.add_argument('--modules', '-m', nargs='*', help='Accessibility modules to apply', 
                   choices=['adhd', 'visual', 'dyslexia', 'autism', 'audio', 'beginner', 'esl'])
parser.add_argument('--subject', '-s', help='Subject area override', 
                   choices=['medical', 'psychology', 'education', 'social_science', 'science'])
parser.add_argument('--output', '-o', help='Output filename')
parser.add_argument('--list-modules', action='store_true', help='List available modules')

args = parser.parse_args()

translator = AcademicTranslator()

if args.list_modules:
    print("📚 Available Accessibility Modules:")
    for module_name in translator.available_modules:
        module = translator.load_module(module_name)
        if module:
            print(f"   • {module.get_name()}: {module.get_description()}")
    return

# Get text input
text = ""
source_file = ""

if args.file:
    print(f"📄 Reading academic document: {args.file}")
    text = translator.extract_text_from_file(args.file)
    source_file = args.file
elif args.text:
    text = args.text
    source_file = "Direct input"
else:
    print("❌ Please provide --file or --text")
    return

if not text or len(text.strip()) < 200:
    print("❌ Not enough text to translate (need at least 200 characters)")
    return

print(f"🧠 Analyzing academic content ({len(text):,} characters)...")

# Apply modules
modules = args.modules or []
if modules:
    print(f"🔧 Applying accessibility modules: {', '.join(modules)}")

# Translate
result = translator.translate_academic_document(
    text, 
    modules=modules,
    source_file=source_file
)

# Show results
print(f"\n📚 Subject Area: {result.subject_area.title()}")
print(f"📖 Reading Level: {result.reading_level}")
print(f"🎯 Translation Confidence: {result.confidence_score:.0%}")

if result.key_findings:
    print(f"\n🔍 Key Findings ({len(result.key_findings)} found):")
    for finding in result.key_findings[:3]:
        print(f"   {finding}")

if result.why_this_matters:
    print(f"\n💡 Why This Matters ({len(result.why_this_matters)} reasons):")
    for matter in result.why_this_matters[:3]:
        print(f"   {matter}")

if result.questions_to_ask:
    print(f"\n❓ Questions to Ask Experts ({len(result.questions_to_ask)} suggestions):")
    for question in result.questions_to_ask[:2]:
        print(f"   {question}")

# Save detailed report
output_name = args.output or f"{result.subject_area}_research_translation"
output_name = re.sub(r'[^a-zA-Z0-9_]', '_', output_name)

translator.save_translation(result, output_name)
print(f"\n💾 Full translation saved: academic_translations/{output_name}.html")

if modules:
    print(f"🔧 Modules applied: {', '.join(result.modules_applied)}")

print(f"\n✨ This research is now accessible to {result.reading_level.lower()} readers!")
```

if **name** == “**main**”:
main()
